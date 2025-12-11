import logging
import os
import uuid
import tempfile
import mimetypes
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Request, status, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel

from open_webui.utils.auth import get_verified_user
from open_webui.utils.access_control import has_permission
from open_webui.utils import libretranslate
from open_webui.constants import ERROR_MESSAGES

log = logging.getLogger(__name__)

router = APIRouter()

# Temporary directory for uploaded files
TEMP_DIR = tempfile.gettempdir()
TRANSLATION_TEMP_DIR = os.path.join(TEMP_DIR, "open_webui_translations")
os.makedirs(TRANSLATION_TEMP_DIR, exist_ok=True)

# File upload limits
MAX_FILE_SIZE = int(os.environ.get("TRANSLATION_FILE_MAX_SIZE", 5242880))  # 5MB default
ALLOWED_FILE_EXTENSIONS = [".pdf", ".txt", ".doc", ".docx"]


def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """
    Extract text content from uploaded file and convert to markdown-like format.
    """
    try:
        if file_ext == ".txt":
            # Plain text - read directly
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif file_ext == ".pdf":
            # PDF extraction using PyPDF2
            try:
                import PyPDF2
                text_content = []
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"## Page {page_num}\n\n{text}\n")
                return "\n".join(text_content)
            except ImportError:
                raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")

        elif file_ext in [".doc", ".docx"]:
            # Word document extraction using python-docx
            try:
                import docx
                doc = docx.Document(file_path)
                text_content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Preserve heading styles
                        if para.style.name.startswith('Heading'):
                            level = para.style.name.replace('Heading ', '')
                            if level.isdigit():
                                text_content.append(f"{'#' * int(level)} {para.text}\n")
                            else:
                                text_content.append(f"## {para.text}\n")
                        else:
                            text_content.append(para.text + "\n")

                # Extract tables
                for table in doc.tables:
                    table_md = "\n"
                    for i, row in enumerate(table.rows):
                        cells = [cell.text.strip() for cell in row.cells]
                        table_md += "| " + " | ".join(cells) + " |\n"
                        if i == 0:  # Add header separator
                            table_md += "| " + " | ".join(["---"] * len(cells)) + " |\n"
                    text_content.append(table_md)

                return "\n".join(text_content)
            except ImportError:
                raise Exception("python-docx not installed. Install with: pip install python-docx")

        else:
            raise Exception(f"Unsupported file type: {file_ext}")

    except Exception as e:
        log.error(f"Text extraction error: {e}")
        raise


def translate_docx_preserving_format(
    file_path: str,
    source_lang: str,
    target_lang: str
) -> str:
    """
    Translate a DOCX file while preserving all formatting, styles, and structure.
    Returns path to the translated file.
    """
    try:
        import docx
        from docx.shared import RGBColor

        # Open the document
        doc = docx.Document(file_path)

        # Translate paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Get the original text
                original_text = paragraph.text

                # Translate the text
                try:
                    translation_result = libretranslate.translate_text(
                        original_text,
                        source_lang,
                        target_lang
                    )
                    translated_text = translation_result['translatedText']

                    # Clear existing runs and add translated text preserving formatting of first run
                    if paragraph.runs:
                        # Preserve formatting from the first run
                        first_run = paragraph.runs[0]
                        font_name = first_run.font.name
                        font_size = first_run.font.size
                        bold = first_run.font.bold
                        italic = first_run.font.italic
                        underline = first_run.font.underline
                        font_color = first_run.font.color.rgb if first_run.font.color.rgb else None

                        # Clear all runs
                        for run in paragraph.runs:
                            run.text = ''

                        # Add translated text with preserved formatting
                        new_run = paragraph.runs[0]
                        new_run.text = translated_text
                        new_run.font.name = font_name
                        new_run.font.size = font_size
                        new_run.font.bold = bold
                        new_run.font.italic = italic
                        new_run.font.underline = underline
                        if font_color:
                            new_run.font.color.rgb = font_color
                    else:
                        # No runs, just set text
                        paragraph.text = translated_text

                except Exception as e:
                    log.warning(f"Failed to translate paragraph: {e}")
                    continue

        # Translate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            original_text = paragraph.text

                            try:
                                translation_result = libretranslate.translate_text(
                                    original_text,
                                    source_lang,
                                    target_lang
                                )
                                translated_text = translation_result['translatedText']

                                # Preserve formatting
                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    font_name = first_run.font.name
                                    font_size = first_run.font.size
                                    bold = first_run.font.bold
                                    italic = first_run.font.italic

                                    for run in paragraph.runs:
                                        run.text = ''

                                    new_run = paragraph.runs[0]
                                    new_run.text = translated_text
                                    new_run.font.name = font_name
                                    new_run.font.size = font_size
                                    new_run.font.bold = bold
                                    new_run.font.italic = italic
                                else:
                                    paragraph.text = translated_text

                            except Exception as e:
                                log.warning(f"Failed to translate table cell: {e}")
                                continue

        # Translate headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    try:
                        translation_result = libretranslate.translate_text(
                            paragraph.text,
                            source_lang,
                            target_lang
                        )
                        translated_text = translation_result['translatedText']

                        if paragraph.runs:
                            for run in paragraph.runs:
                                run.text = ''
                            paragraph.runs[0].text = translated_text
                        else:
                            paragraph.text = translated_text
                    except Exception as e:
                        log.warning(f"Failed to translate header: {e}")

            # Translate footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    try:
                        translation_result = libretranslate.translate_text(
                            paragraph.text,
                            source_lang,
                            target_lang
                        )
                        translated_text = translation_result['translatedText']

                        if paragraph.runs:
                            for run in paragraph.runs:
                                run.text = ''
                            paragraph.runs[0].text = translated_text
                        else:
                            paragraph.text = translated_text
                    except Exception as e:
                        log.warning(f"Failed to translate footer: {e}")

        # Save the translated document
        translated_file_path = file_path.replace('.docx', '_translated.docx')
        doc.save(translated_file_path)

        log.info(f"DOCX translated with format preservation: {translated_file_path}")
        return translated_file_path

    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        log.error(f"DOCX translation error: {e}")
        raise


class LanguageResponse(BaseModel):
    code: str
    name: str


class LanguagesResponse(BaseModel):
    languages: list[LanguageResponse]


class DetectLanguageRequest(BaseModel):
    text: str


class DetectLanguageResponse(BaseModel):
    language: str
    confidence: float


class TranslateRequest(BaseModel):
    text: str
    source: str
    target: str


class TranslateResponse(BaseModel):
    translatedText: str
    detectedLanguage: Optional[str] = None


class FileUploadResponse(BaseModel):
    fileId: str
    filename: str
    size: int
    extractedText: str
    fileExtension: str


class TranslateFileRequest(BaseModel):
    fileId: str
    source: str
    target: str


class TranslateFileResponse(BaseModel):
    translatedText: Optional[str] = None  # For TXT/PDF (text-only formats)
    translatedFileId: Optional[str] = None  # For DOCX (format-preserving)
    detectedLanguage: Optional[str] = None
    filename: str
    preservedFormat: bool  # True if format was preserved (DOCX)


############################
# Get Available Languages
############################


@router.get("/languages", response_model=LanguagesResponse)
async def get_languages(
    request: Request,
    user=Depends(get_verified_user)
):
    """
    Get list of available languages from LibreTranslate.
    """
    if user.role != "admin" and not has_permission(
        user.id, "features.translation", request.app.state.config.USER_PERMISSIONS
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    try:
        languages = libretranslate.get_languages()
        return LanguagesResponse(
            languages=[LanguageResponse(**lang) for lang in languages]
        )
    except libretranslate.LibreTranslateError as e:
        log.error(f"LibreTranslate error: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(e)
        )
    except Exception as e:
        log.exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=ERROR_MESSAGES.DEFAULT()
        )


############################
# Detect Language
############################


@router.post("/detect", response_model=DetectLanguageResponse)
async def detect_language(
    request: Request,
    form_data: DetectLanguageRequest,
    user=Depends(get_verified_user)
):
    """
    Auto-detect the language of provided text.
    """
    if user.role != "admin" and not has_permission(
        user.id, "features.translation", request.app.state.config.USER_PERMISSIONS
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    try:
        result = libretranslate.detect_language(form_data.text)
        return DetectLanguageResponse(**result)
    except libretranslate.LibreTranslateError as e:
        log.error(f"LibreTranslate error: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(e)
        )
    except Exception as e:
        log.exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=ERROR_MESSAGES.DEFAULT()
        )


############################
# Translate Text
############################


@router.post("/translate", response_model=TranslateResponse)
async def translate_text(
    request: Request,
    form_data: TranslateRequest,
    user=Depends(get_verified_user)
):
    """
    Translate text from source language to target language.
    """
    if user.role != "admin" and not has_permission(
        user.id, "features.translation", request.app.state.config.USER_PERMISSIONS
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    try:
        result = libretranslate.translate_text(
            form_data.text,
            form_data.source,
            form_data.target
        )
        return TranslateResponse(**result)
    except libretranslate.LibreTranslateError as e:
        log.error(f"LibreTranslate error: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(e)
        )
    except Exception as e:
        log.exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=ERROR_MESSAGES.DEFAULT()
        )


############################
# Upload File for Translation
############################


@router.post("/file/upload", response_model=FileUploadResponse)
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    user=Depends(get_verified_user)
):
    """
    Upload a file for translation.
    Validates file type and size, extracts text content to markdown.
    """
    if user.role != "admin" and not has_permission(
        user.id, "features.translation", request.app.state.config.USER_PERMISSIONS
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    # Validate file extension
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_FILE_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type. Allowed types: {', '.join(ALLOWED_FILE_EXTENSIONS)}"
        )

    # Read file to check size
    file_content = await file.read()
    file_size = len(file_content)

    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds maximum allowed size of {MAX_FILE_SIZE / 1024 / 1024:.1f}MB"
        )

    if file_size == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File is empty"
        )

    try:
        # Generate unique file ID
        file_id = str(uuid.uuid4())
        file_path = os.path.join(TRANSLATION_TEMP_DIR, f"{file_id}{file_ext}")

        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)

        # Extract text content based on file type
        try:
            extracted_text = extract_text_from_file(file_path, file_ext)
        except Exception as e:
            log.error(f"Text extraction failed: {e}")
            # Clean up the file
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to extract text from file: {str(e)}"
            )

        log.info(f"File uploaded and text extracted: {file_id} ({file.filename}, {file_size} bytes, {len(extracted_text)} chars)")

        return FileUploadResponse(
            fileId=file_id,
            filename=file.filename,
            size=file_size,
            extractedText=extracted_text,
            fileExtension=file_ext
        )

    except HTTPException:
        raise
    except Exception as e:
        log.exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to save uploaded file"
        )


############################
# Translate File
############################


@router.post("/file/translate", response_model=TranslateFileResponse)
async def translate_file(
    request: Request,
    form_data: TranslateFileRequest,
    user=Depends(get_verified_user)
):
    """
    Translate an uploaded file.
    For DOCX: Preserves formatting and returns file ID for download.
    For TXT/PDF: Extracts text, translates, and returns text.
    """
    if user.role != "admin" and not has_permission(
        user.id, "features.translation", request.app.state.config.USER_PERMISSIONS
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    # Find the uploaded file
    matching_files = [
        f for f in os.listdir(TRANSLATION_TEMP_DIR)
        if f.startswith(form_data.fileId)
    ]

    if not matching_files:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found. Please upload the file again."
        )

    source_file_path = os.path.join(TRANSLATION_TEMP_DIR, matching_files[0])
    file_ext = os.path.splitext(matching_files[0])[1].lower()

    try:
        # DOCX: Use format-preserving translation
        if file_ext == ".docx":
            log.info(f"Using format-preserving translation for DOCX: {form_data.fileId}")

            translated_file_path = translate_docx_preserving_format(
                source_file_path,
                form_data.source,
                form_data.target
            )

            # Generate a unique ID for the translated file
            translated_file_id = str(uuid.uuid4())
            translated_file_name = os.path.basename(translated_file_path)

            # Rename the translated file with the new ID
            final_translated_path = os.path.join(
                TRANSLATION_TEMP_DIR,
                f"{translated_file_id}_translated.docx"
            )
            os.rename(translated_file_path, final_translated_path)

            # Get original filename without the UUID
            original_filename = matching_files[0].replace(form_data.fileId, "").lstrip("_")
            if not original_filename:
                original_filename = "document.docx"

            log.info(f"DOCX file translated with format preservation: {translated_file_id}")

            return TranslateFileResponse(
                translatedText=None,
                translatedFileId=translated_file_id,
                detectedLanguage=None,
                filename=original_filename,
                preservedFormat=True
            )

        # TXT/PDF: Extract text and translate
        else:
            log.info(f"Using text-based translation for {file_ext}: {form_data.fileId}")

            # Extract text from file
            extracted_text = extract_text_from_file(source_file_path, file_ext)

            if not extracted_text or not extracted_text.strip():
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="No text content found in file"
                )

            # Translate the extracted text
            translation_result = libretranslate.translate_text(
                extracted_text,
                form_data.source,
                form_data.target
            )

            log.info(f"File translated: {form_data.fileId} ({len(extracted_text)} chars -> {len(translation_result['translatedText'])} chars)")

            # Extract original filename
            original_filename = matching_files[0].replace(form_data.fileId, "").lstrip("_").lstrip(file_ext) or "document"

            return TranslateFileResponse(
                translatedText=translation_result['translatedText'],
                translatedFileId=None,
                detectedLanguage=translation_result.get('detectedLanguage'),
                filename=original_filename,
                preservedFormat=False
            )

    except HTTPException:
        raise
    except libretranslate.LibreTranslateError as e:
        log.error(f"LibreTranslate error: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(e)
        )
    except Exception as e:
        log.exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"File translation failed: {str(e)}"
        )


############################
# Download Translated File
############################


@router.get("/file/download/{file_id}")
async def download_translated_file(
    request: Request,
    file_id: str,
    user=Depends(get_verified_user)
):
    """
    Download a translated DOCX file.
    """
    if user.role != "admin" and not has_permission(
        user.id, "features.translation", request.app.state.config.USER_PERMISSIONS
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    # Find the translated file
    matching_files = [
        f for f in os.listdir(TRANSLATION_TEMP_DIR)
        if f.startswith(file_id) and f.endswith('.docx')
    ]

    if not matching_files:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Translated file not found"
        )

    file_path = os.path.join(TRANSLATION_TEMP_DIR, matching_files[0])

    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found"
        )

    # Generate download filename
    download_filename = f"translated_{file_id}.docx"

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=download_filename,
        headers={
            "Content-Disposition": f'attachment; filename="{download_filename}"'
        }
    )


############################
# Cleanup Old Files (Helper)
############################


def cleanup_old_translation_files(max_age_hours: int = 1):
    """
    Clean up translation files older than max_age_hours.
    This should be called periodically by a background task.
    """
    import time

    try:
        current_time = time.time()
        max_age_seconds = max_age_hours * 3600

        for filename in os.listdir(TRANSLATION_TEMP_DIR):
            file_path = os.path.join(TRANSLATION_TEMP_DIR, filename)
            if os.path.isfile(file_path):
                file_age = current_time - os.path.getmtime(file_path)
                if file_age > max_age_seconds:
                    os.remove(file_path)
                    log.info(f"Cleaned up old translation file: {filename}")

    except Exception as e:
        log.error(f"Error cleaning up translation files: {e}")
